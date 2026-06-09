# -*- coding: utf-8 -*-
"""
法典转换器：把 法典源/*.docx 解析成网站用的结构化数据。
输出：
  site/data/<codexId>.json   每本法典的词条 + 目录树
  site/data/codexes.json     法典索引（给顶部切换用）
  site/data/待复核_<codexId>.txt  可能解析有误的词条，供人工复核
用法：python tools/convert.py   （或双击 转换法典.bat）
"""
import argparse, os, re, io, json, hashlib, glob, shutil
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.oxml.ns import qn
from PIL import Image

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.join(ROOT, "法典源")
DATA_DIR = os.path.join(ROOT, "site", "data")
IMG_DIR = os.path.join(ROOT, "site", "images")
ORIG_DIR = os.path.join(ROOT, "originals")
ARCHIVE_DIR = os.path.join(ROOT, "法典源_已转换")
ARCHIVE_MANIFEST = os.path.join(ARCHIVE_DIR, "processed_sources.json")
os.makedirs(DATA_DIR, exist_ok=True)

# 已知法典 → 固定 id（其余用文件名哈希）
ID_MAP = [
    ("所长常规", "suozhang"),   # 原「所长常规」法典固定用 suozhang（已配图，勿改）
    ("所长色色NovalAI个人法典（上）", "codex_6e699406"),
    ("所长色色NovalAI个人法典(上)", "codex_6e699406"),
    ("所长色色NovelAI个人法典（上）", "codex_6e699406"),
    ("所长色色NovelAI个人法典(上)", "codex_6e699406"),
    ("所长色色NovalAI个人法典（下）", "codex_8489ac52"),
    ("所长色色NovalAI个人法典(下)", "codex_8489ac52"),
    ("所长色色NovelAI个人法典（下）", "codex_8489ac52"),
    ("所长色色NovelAI个人法典(下)", "codex_8489ac52"),
]
META_OVERRIDES = {
    "suozhang": {"author": "戒红所"},
    "codex_6e699406": {"title": "所长色色NovalAI个人法典(上)", "author": "戒红所"},
    "codex_8489ac52": {"title": "所长色色NovalAI个人法典(下)", "author": "戒红所"},
    "codex_3f60585d": {"title": "涩涩法典(梦神版)", "author": "梦神"},
}
IMG_EXTS = ["jpg", "jpeg", "png", "webp", "gif", "avif"]
MAXDIM = 1100

def normalized_stem(stem):
    return (
        re.sub(r"\s+", "", stem)
        .replace("（", "(")
        .replace("）", ")")
        .lower()
        .replace("novelai", "novalai")
    )

def codex_id(stem):
    norm = normalized_stem(stem)
    for key, cid in ID_MAP:
        if normalized_stem(key) in norm:
            return cid
    if "所长色色" in stem:
        if "(上)" in norm:
            return "codex_6e699406"
        if "(下)" in norm:
            return "codex_8489ac52"
    return "codex_" + hashlib.md5(stem.encode("utf-8")).hexdigest()[:8]

def parse_meta(stem):
    ver = re.search(r"(\d{4}[.\-]\d{1,2}[.\-]?\d{0,2})", stem)
    title = re.sub(r"[（(]\d{4}[.\-]\d{1,2}.*$", "", stem).strip()
    if title == stem:
        title = re.split(r"[（(]", stem)[0].strip()
    author = ""
    m = re.search(r"[（(].*?([一-鿿]{2,})整理", stem)
    if m:
        author = m.group(1)
    return title or stem, (ver.group(1) if ver else ""), author

def is_cjk(c):
    return "一" <= c <= "鿿"

def has_cjk(t):
    return any(is_cjk(c) for c in t)

def cjk_ratio(t):
    cjk = sum(1 for c in t if is_cjk(c))
    al = sum(1 for c in t if c.isalnum() or is_cjk(c))
    return cjk / al if al else 0.0

def classify(t):
    """把一段文本判定为 词条标题(title) / tag行(tag) / 其它(note)"""
    r = cjk_ratio(t)
    if r > 0.5 and len(t) < 45:          # 中文为主的短行 → 标题（即便夹少量英文/括号）
        return "title"
    if re.match(r"^(角色|人物)\d*[：:]", t):
        return "tag"
    if re.match(r"^[一-鿿]", t):
        return "title"                   # 中文开头的标题常会夹英文 tag 注释
    has_latin = bool(re.search(r"[A-Za-z]", t))
    has_tagsig = any(s in t for s in [",", "，", "::", "{", "}", "[", "]", "_"])
    if has_latin and has_tagsig:          # 含英文且有tag语法 → tag行
        return "tag"
    if len(t) < 45 and has_cjk(t):        # 中英混排的短标题（如 上古galgame风格）
        return "title"
    if len(t) < 35 and has_latin:         # 极短的纯英文标签 → 当标题
        return "title"
    return "note"

def visible_lines(t):
    """Word 段落里可能用软换行塞了多条内容；按肉眼看到的行拆开解析。"""
    return [x.strip() for x in t.splitlines() if x.strip()]

def paragraph_text(p):
    """python-docx 的 p.text 不包含文本框；新版梦神法典把 tag 放在文本框里。"""
    text = "".join(t.text or "" for t in p._p.iter() if t.tag.endswith("}t"))
    return text.strip()

def paragraph_blips(p):
    return [b for b in p._p.xpath('.//*[local-name()="blip"]')]

def paragraph_image_parts(doc, p):
    parts = []
    seen = set()
    for blip in paragraph_blips(p):
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        part = doc.part.related_parts.get(rid)
        blob = getattr(part, "blob", None) if part is not None else None
        if blob:
            sig = hashlib.sha1(blob).hexdigest()
            if sig in seen:
                continue
            seen.add(sig)
            parts.append(part)
    return parts

def is_dictionary_path(path):
    return len(path) >= 2 and path[0] == "各式场景" and path[1] == "视角与打光"

def should_skip_path(path):
    return path[:1] == ["自然语言"]

def dictionary_term(t):
    if "：" in t:
        left, right = t.split("：", 1)
    elif ":" in t:
        left, right = t.split(":", 1)
    else:
        return None
    term = left.strip()
    desc = right.strip()
    if not term or not desc:
        return None
    if not re.search(r"[A-Za-z]", term):
        return None
    if not has_cjk(desc):
        return None
    if len(term) > 80:
        return None
    return term

def strip_number_prefix(t):
    m = re.match(r"^\s*(\d+)[，,、.．]\s*(.+)$", t)
    if m:
        return m.group(1), m.group(2).strip()
    return None, t.strip()

def is_artist_group_entry(e):
    return e["path"][-1:] == ["编纂者常用画师组"] and e["title"].startswith("NAI")

def expand_special_entries(entries):
    """把已知的复合词条拆成更适合网页复制的独立卡片。"""
    out = []
    for e in entries:
        if not is_artist_group_entry(e):
            out.append(e)
            continue

        tag_lines = []
        for block in e["tags"]:
            tag_lines.extend(visible_lines(block))

        base = e["title"].rstrip("：:")
        for i, line in enumerate(tag_lines, 1):
            num, tags = strip_number_prefix(line)
            label = num or str(i)
            title = e["title"] if i == 1 else f"{base}：{label}"
            out.append({
                "title": title,
                "path": e["path"],
                "tags": [tags],
                "isNew": e["isNew"],
            })
    return out

def load_existing_entries(cid):
    jp = os.path.join(DATA_DIR, cid + ".json")
    if not os.path.exists(jp):
        return []
    try:
        with open(jp, encoding="utf-8") as f:
            data = json.load(f)
        return data.get("entries", [])
    except Exception:
        return []

def norm_tags(t):
    return re.sub(r"\s+", " ", (t or "").replace("\u00a0", " ")).strip()

def tag_match_score(new_tags, old_tags):
    new = norm_tags(new_tags)
    old = norm_tags(old_tags)
    if not new or not old:
        return 0
    if new == old:
        return 1000000
    if old.startswith(new):
        return 800000 + len(new)
    if new.startswith(old):
        return 700000 + len(old)

    n = 0
    for a, b in zip(new, old):
        if a != b:
            break
        n += 1
    return n

def assign_stable_ids(cid, items):
    """复用旧 JSON 中的 id，避免修解析规则后把已配图词条整体错位。"""
    old_entries = load_existing_entries(cid)
    if not old_entries:
        final = []
        for i, item in enumerate(items, 1):
            eid = f"{cid}-{i:04d}"
            final.append({**item, "id": eid, **image_metadata(cid, eid)})
        return final

    old_by_key = defaultdict(list)
    old_ids = set()
    max_n = 0
    id_re = re.compile(r"^" + re.escape(cid) + r"-(\d+)$")
    for old in old_entries:
        oid = old.get("id")
        if oid:
            old_ids.add(oid)
            m = id_re.match(oid)
            if m:
                max_n = max(max_n, int(m.group(1)))
        key = (tuple(old.get("path", [])), old.get("title", ""))
        old_by_key[key].append(old)

    used = set()
    next_n = max_n + 1

    def fresh_id():
        nonlocal next_n
        while True:
            eid = f"{cid}-{next_n:04d}"
            next_n += 1
            if eid not in used and eid not in old_ids:
                return eid

    final = []
    for item in items:
        key = (tuple(item["path"]), item["title"])
        best = None
        best_score = -1
        for old in old_by_key.get(key, []):
            oid = old.get("id")
            if not oid or oid in used:
                continue
            score = tag_match_score(item["tags"], old.get("tags", ""))
            if score > best_score:
                best = old
                best_score = score

        eid = best.get("id") if best is not None else fresh_id()
        used.add(eid)
        final.append({**item, "id": eid, **image_metadata(cid, eid, best)})
    return final

def outline_lvl(p):
    pPr = p._p.pPr
    if pPr is None:
        return None
    o = pPr.find(qn("w:outlineLvl"))
    return int(o.get(qn("w:val"))) if o is not None else None

def outline_lvl_with_style(p):
    lv = outline_lvl(p)
    if lv is not None:
        return lv
    try:
        pPr = p.style.element.pPr
        if pPr is not None:
            o = pPr.find(qn("w:outlineLvl"))
            if o is not None:
                return int(o.get(qn("w:val")))
    except Exception:
        pass
    return None

def is_pink(p):
    for r in p.runs:
        rPr = r._element.rPr
        if rPr is not None:
            hh = rPr.find(qn("w:highlight"))
            if hh is not None and hh.get(qn("w:val")) not in (None, "none"):
                return True
    return False

def find_image(cid, eid):
    for ext in IMG_EXTS:
        fn = os.path.join(IMG_DIR, cid, eid + "." + ext)
        if os.path.exists(fn):
            return eid + "." + ext
    return None

def find_original(cid, eid):
    for ext in IMG_EXTS:
        fn = os.path.join(ORIG_DIR, cid, eid + "." + ext)
        if os.path.exists(fn):
            return eid + "." + ext
    return None

def local_asset_rev(cid, image, original):
    h = hashlib.sha256()
    found = False
    for root, fn in ((IMG_DIR, image), (ORIG_DIR, original)):
        if not fn:
            continue
        path = os.path.join(root, cid, fn)
        if not os.path.exists(path):
            continue
        st = os.stat(path)
        h.update(f"{fn}:{st.st_size}:{st.st_mtime_ns}".encode("utf-8"))
        found = True
    return h.hexdigest()[:16] if found else None

def image_dimensions(cid, image, old=None):
    old = old or {}
    if image:
        path = os.path.join(IMG_DIR, cid, image)
        if os.path.exists(path):
            try:
                with Image.open(path) as im:
                    return im.size
            except Exception:
                pass
    w = old.get("imageWidth")
    h = old.get("imageHeight")
    return (w, h) if w and h else None

def image_metadata(cid, eid, old=None):
    old = old or {}
    image = find_image(cid, eid) or old.get("image")
    original = find_original(cid, eid) or old.get("original")
    asset_rev = local_asset_rev(cid, image, original) or old.get("assetRev")
    meta = {"image": image}
    dims = image_dimensions(cid, image, old)
    if dims:
        meta["imageWidth"], meta["imageHeight"] = dims
    if original:
        meta["original"] = original
    if asset_rev:
        meta["assetRev"] = asset_rev
    return meta

def main_tag_chunks(text):
    """新版梦神法典：从文本框里的「主要Tag：...」抽出正向主 tag，去掉备注/负面。"""
    text = re.sub(r"\s+", " ", text or "").strip()
    out = []
    pat = r"主要\s*[Tt]ag\s*[：:]"
    for m in re.finditer(pat, text):
        start = m.end()
        nxt = re.search(pat, text[start:])
        stop = len(text) if not nxt else start + nxt.start()
        chunk = text[start:stop]
        chunk = clean_mengshen_tag_chunk(chunk)
        if chunk and chunk not in out:
            out.append(chunk)
    return out

def clean_mengshen_tag_chunk(chunk):
    chunk = re.sub(r"\s+", " ", chunk or "").strip()
    cut = re.search(r"(?:负面\s*(?:[Tt]ag|tag)?|备注)\s*[：:]?", chunk)
    if cut:
        chunk = chunk[:cut.start()]
    return chunk.strip(" ：:，,。\\")

ROLE_GUIDE_RE = re.compile(r"\u89d2\u8272\u680f\u5199\u6cd5\s*[\uff1a:]")
MAIN_TAG_LABEL_RE = re.compile(r"\u4e3b\s*[Tt]ag\s*[\uff1a:]", re.I)
ROLE_LABEL_RE = re.compile(r"\u89d2\u8272\s*\d*\s*[\uff1a:]")

def normalize_mengshen_tag_body(body):
    body = clean_mengshen_tag_chunk(body)
    guide = ROLE_GUIDE_RE.search(body)
    if guide:
        prefix = body[:guide.start()]
        body = prefix if re.search(r"[A-Za-z]", prefix) else body[guide.end():]
    body = MAIN_TAG_LABEL_RE.sub("", body)
    body = ROLE_LABEL_RE.sub(",", body)
    body = re.sub(r"\s+", " ", body)
    body = re.sub(r",\s*,+", ",", body)
    return body.strip(" ：:，,。\\")

VERSION_LABEL_RE = re.compile(
    r"(?P<label>(?:[\u4e00-\u9fff][\u4e00-\u9fffA-Za-z0-9.\uff08\uff09()&/]*?)?"
    r"(?:\u7248[Tt]ag|\u7248\u672c|\u820c\u5934|\u7248|\u4f4d|\u8138|\u9762)"
    r"(?:\s*(?:\([^)]+\)|\uff08[^\uff09]+\uff09))?)"
    r"\s*[\uff1a:]?(?=\s*(?:[A-Za-z0-9{_\-]|\u4e3b\s*[Tt]ag|$))"
)

def version_label_spans(chunk):
    spans = []
    for m in VERSION_LABEL_RE.finditer(chunk):
        start, end = m.span()
        label = m.group("label").strip()
        j = start
        while j > 0 and re.match(r"[A-Za-z0-9.]", chunk[j - 1]):
            j -= 1
        prefix = chunk[j:start]
        prefix_match = re.search(r"((?:nai)?\d+(?:\.\d+)?|sex)$", prefix, re.I)
        if prefix_match:
            prefix = prefix_match.group(1)
            start = start - len(prefix)
            label = prefix + label
        spans.append((start, end, label))
    return spans

def split_versioned_tag_chunk(chunk):
    """把「无衣版：tags 着衣版：tags」拆成多个变体。"""
    chunk = clean_mengshen_tag_chunk(chunk)
    if not chunk:
        return []
    spans = version_label_spans(chunk)
    if not spans:
        body = normalize_mengshen_tag_body(chunk)
        return [(None, body)] if body else []

    variants = []
    if spans[0][0] > 0:
        head = clean_mengshen_tag_chunk(chunk[:spans[0][0]])
        if head:
            head = normalize_mengshen_tag_body(head)
            if head:
                variants.append((None, head))

    for i, (_, label_end, label) in enumerate(spans):
        start = label_end
        end = spans[i + 1][0] if i + 1 < len(spans) else len(chunk)
        body = normalize_mengshen_tag_body(chunk[start:end])
        if body:
            variants.append((label, body))

    # 同一个文本框在 docx XML 中偶尔会重复一遍，按 tag 去重。
    out, seen = [], set()
    for label, body in variants:
        key = norm_tags(body)
        if key in seen:
            continue
        seen.add(key)
        out.append((label, body))
    return out

CJK_TAG_RE = re.compile(r"[\u4e00-\u9fff]")
MENGSHEN_ALLOWED_CJK_TAG_TOKENS = {
    "场景",
    "武器",
    "下身服装",
    "动作",
    "姿势",
    "内容",
    "食物",
}

def unknown_mengshen_cjk_tokens(tags):
    return [
        token for token in re.findall(r"[\u4e00-\u9fff]+", tags)
        if token not in MENGSHEN_ALLOWED_CJK_TAG_TOKENS
    ]

def should_review_mengshen_tags(item):
    if not CJK_TAG_RE.search(item["tags"]):
        return False
    return bool(unknown_mengshen_cjk_tokens(item["tags"]))

def embedded_image_ext(part):
    ext = os.path.splitext(str(part.partname))[1].lower().lstrip(".")
    if ext == "jpeg":
        return "jpg"
    if ext in IMG_EXTS:
        return ext
    ctype = getattr(part, "content_type", "")
    if ctype == "image/jpeg":
        return "jpg"
    if ctype.startswith("image/"):
        guess = ctype.split("/", 1)[1].lower()
        return "jpg" if guess == "jpeg" else guess
    return "png"

def save_embedded_image(cid, eid, parts):
    if not parts:
        return
    part = max(parts, key=lambda p: len(p.blob or b""))
    raw = part.blob
    if not raw:
        return

    original_dir = os.path.join(ORIG_DIR, cid)
    thumb_dir = os.path.join(IMG_DIR, cid)
    os.makedirs(original_dir, exist_ok=True)
    os.makedirs(thumb_dir, exist_ok=True)

    ext = embedded_image_ext(part)
    original_path = os.path.join(original_dir, f"{eid}.{ext}")
    thumb_path = os.path.join(thumb_dir, f"{eid}.jpg")

    with open(original_path, "wb") as f:
        f.write(raw)

    image = Image.open(io.BytesIO(raw))
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")
    image.thumbnail((MAXDIM, MAXDIM), Image.LANCZOS)
    image.save(thumb_path, "JPEG", quality=86, optimize=True)

def build_tree(entries):
    root = {}
    for e in entries:
        node = root
        for name in e["path"]:
            node = node.setdefault(name, {"_count": 0, "_children": {}})
            node["_count"] += 1
            node = node["_children"]
    def to_list(d):
        out = []
        for name, v in d.items():
            out.append({"name": name, "count": v["_count"], "children": to_list(v["_children"])})
        return out
    return to_list(root)

def convert_mengshen_docx(path, cid, title, ver, author, doc):
    cats = []
    entries = []
    cur = None

    def add_variants_from_text(entry, text, parts):
        variants = []
        for chunk in main_tag_chunks(text):
            variants.extend(split_versioned_tag_chunk(chunk))
        if not variants:
            return 0

        seen = entry.setdefault("_variant_keys", set())
        for i, (label, tags) in enumerate(variants):
            key = norm_tags(tags)
            if key in seen:
                continue
            seen.add(key)

            if len(variants) > 1:
                image_parts = [parts[i]] if i < len(parts) else []
            else:
                image_parts = [parts[0]] if parts else []

            entry["variants"].append({
                "label": label,
                "tags": tags,
                "image_parts": image_parts,
            })
        return min(len(parts), len(variants)) if len(variants) > 1 else (1 if parts else 0)

    def add_pending_images(entry, parts):
        seen = entry.setdefault("_pending_image_keys", set())
        for part in parts:
            blob = getattr(part, "blob", None)
            if not blob:
                continue
            sig = hashlib.sha1(blob).hexdigest()
            if sig in seen:
                continue
            seen.add(sig)
            entry["pending_images"].append(part)

    def review_item(item, reason):
        copied = dict(item)
        copied["_reviewReason"] = reason
        review.append(copied)

    for p in doc.paragraphs:
        text = paragraph_text(p)
        lv = outline_lvl_with_style(p)
        image_parts = paragraph_image_parts(doc, p)

        if lv is None:
            if cur is None:
                continue
            if image_parts:
                add_pending_images(cur, image_parts)
            used_images = add_variants_from_text(cur, text, cur.get("pending_images", []))
            if main_tag_chunks(text):
                cur["pending_images"] = cur["pending_images"][used_images:]
                cur["_pending_image_keys"] = {
                    hashlib.sha1(getattr(part, "blob", b"")).hexdigest()
                    for part in cur["pending_images"]
                    if getattr(part, "blob", None)
                }
            if is_pink(p):
                cur["isNew"] = True
            continue

        if not text or text == "目录":
            cur = None
            continue

        if lv <= 1:
            while len(cats) <= lv:
                cats.append(None)
            cats[lv] = text
            cats = cats[:lv + 1]
            cur = None
            continue

        if lv == 2:
            cur = {
                "title": text,
                "path": [c for c in cats if c],
                "variants": [],
                "pending_images": [],
                "_pending_image_keys": set(),
                "_variant_keys": set(),
                "isNew": is_pink(p),
            }
            entries.append(cur)
            continue

        cur = None

    items, image_parts_by_key, review = [], {}, []
    for e in entries:
        if not e["variants"]:
            review.append({"title": e["title"], "path": e["path"], "tags": "", "isNew": e["isNew"]})
            continue
        multi = len(e["variants"]) > 1
        for i, variant in enumerate(e["variants"], 1):
            suffix = variant["label"] or f"版本{i}"
            item_title = f"{e['title']}：{suffix}" if multi else e["title"]
            item = {
                "title": item_title,
                "path": e["path"],
                "tags": variant["tags"].strip(),
                "isNew": e["isNew"],
            }
            image_parts_by_key[(tuple(item["path"]), item["title"], item["tags"])] = variant["image_parts"]
            items.append(item)

    final = assign_stable_ids(cid, items)
    for item in final:
        key = (tuple(item["path"]), item["title"], item["tags"])
        parts = image_parts_by_key.get(key, [])
        if parts:
            save_embedded_image(cid, item["id"], parts)
            item.update(image_metadata(cid, item["id"], item))
        else:
            for k in ("image", "imageWidth", "imageHeight", "original", "assetRev"):
                item.pop(k, None)
            item["image"] = None
            review_item(item, "missing image")
        if should_review_mengshen_tags(item):
            review_item(item, "tag contains unknown Chinese")
        if not item.get("image"):
            pass

    tree = build_tree(final)
    imaged = sum(1 for e in final if e["image"])

    with io.open(os.path.join(DATA_DIR, cid + ".json"), "w", encoding="utf-8") as f:
        json.dump({
            "id": cid, "title": title, "version": ver, "author": author,
            "entryCount": len(final), "imagedCount": imaged,
            "tree": tree, "entries": final,
        }, f, ensure_ascii=False)

    review_path = os.path.join(DATA_DIR, f"待复核_{cid}.txt")
    if review:
        with io.open(review_path, "w", encoding="utf-8") as f:
            f.write(f"# {title}：{len(review)} 条可能解析有误，请人工瞄一眼\n\n")
            for e in review:
                reason = e.get("_reviewReason")
                prefix = f"[{reason}] " if reason else ""
                f.write(prefix + " > ".join(e.get("path", [])) + " › " + e["title"] + "\n")
                if e.get("tags"):
                    f.write(e["tags"][:300] + "\n")
                f.write("\n")
    elif os.path.exists(review_path):
        os.remove(review_path)

    return {
        "id": cid, "title": title, "version": ver, "author": author,
        "entryCount": len(final), "imagedCount": imaged, "reviewCount": len(review)}

def convert(path, cid):
    stem = os.path.splitext(os.path.basename(path))[0]
    title, ver, author = parse_meta(stem)
    meta = META_OVERRIDES.get(cid, {})
    title = meta.get("title", title)
    author = meta.get("author", author)
    doc = Document(path)

    if "梦神" in stem and "涩涩法典" in stem:
        return convert_mengshen_docx(path, cid, title, ver, author, doc)

    cats = [None, None, None, None]
    entries, cur = [], None
    seen_toc = False
    for p in doc.paragraphs:
        if p.style.name.startswith("toc"):
            seen_toc = True
            continue
        lines = visible_lines(p.text)
        if not lines:
            continue
        lv = outline_lvl(p)
        if lv is not None and lv <= 3:
            cats[lv] = " ".join(lines)
            for k in range(lv + 1, 4):
                cats[k] = None
            cur = None
            continue
        if not seen_toc:            # 跳过目录之前的零碎
            continue
        path_now = [c for c in cats if c]
        if should_skip_path(path_now):
            cur = None
            continue
        for t in lines:
            term = dictionary_term(t) if is_dictionary_path(path_now) else None
            if term is not None:
                cur = {"title": t, "path": path_now, "tags": [term], "isNew": is_pink(p)}
                entries.append(cur)
                cur = None
                continue

            kind = classify(t)
            if kind == "title":
                cur = {"title": t, "path": path_now, "tags": [], "isNew": is_pink(p)}
                entries.append(cur)
            elif kind == "tag":
                if cur is None:
                    cur = {"title": "(未命名)", "path": path_now, "tags": [], "isNew": False}
                    entries.append(cur)
                cur["tags"].append(t)
                if is_pink(p):
                    cur["isNew"] = True
            # note: 丢弃

    # 仅保留有 tag 的词条；合并 tag 行；复用旧 id；探测配图
    items, review = [], []
    for e in expand_special_entries(entries):
        if not e["tags"]:
            continue
        block = "\n".join(e["tags"]).strip()
        items.append({
            "title": e["title"],
            "path": e["path"],
            "tags": block,
            "isNew": e["isNew"],
        })

    final = assign_stable_ids(cid, items)
    for item in final:
        block = item["tags"]
        is_dict = is_dictionary_path(item["path"])
        # 可疑：tag 块里没有任何逗号/:: → 可能是误判
        if (not is_dict) and ("," not in block) and ("，" not in block) and ("::" not in block):
            review.append(item)
        elif (not is_dict) and len(item["title"]) > 30:
            review.append(item)

    tree = build_tree(final)
    imaged = sum(1 for e in final if e["image"])

    with io.open(os.path.join(DATA_DIR, cid + ".json"), "w", encoding="utf-8") as f:
        json.dump({
            "id": cid, "title": title, "version": ver, "author": author,
            "entryCount": len(final), "imagedCount": imaged,
            "tree": tree, "entries": final,
        }, f, ensure_ascii=False)

    if review:
        with io.open(os.path.join(DATA_DIR, f"待复核_{cid}.txt"), "w", encoding="utf-8") as f:
            f.write(f"# {title}：{len(review)} 条可能解析有误，请人工瞄一眼\n\n")
            for e in review:
                f.write(f"[{' › '.join(e['path'])}] {e['title']}\n    {e['tags'][:120]}\n\n")

    return {"id": cid, "title": title, "version": ver, "author": author,
            "entryCount": len(final), "imagedCount": imaged, "reviewCount": len(review)}

def load_archived_source_names():
    if not os.path.exists(ARCHIVE_MANIFEST):
        return set()
    try:
        with open(ARCHIVE_MANIFEST, encoding="utf-8") as f:
            data = json.load(f)
        return set(data.get("sourceNames", []))
    except Exception:
        return set()

def write_archived_source_names(names):
    os.makedirs(ARCHIVE_DIR, exist_ok=True)
    with open(ARCHIVE_MANIFEST, "w", encoding="utf-8") as f:
        json.dump({"sourceNames": sorted(names)}, f, ensure_ascii=False, indent=2)

def archive_sources(docs):
    docs = [d for d in docs if os.path.exists(d) and not os.path.basename(d).startswith("~$")]
    if not docs:
        return {"moved": [], "copied": [], "failed": []}
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    dest_dir = os.path.join(ARCHIVE_DIR, stamp)
    os.makedirs(dest_dir, exist_ok=True)
    archived_names = load_archived_source_names()
    result = {"moved": [], "copied": [], "failed": []}
    for src in docs:
        name = os.path.basename(src)
        dest = os.path.join(dest_dir, name)
        if os.path.exists(dest):
            base, ext = os.path.splitext(name)
            n = 2
            while os.path.exists(dest):
                dest = os.path.join(dest_dir, f"{base}_{n}{ext}")
                n += 1
        try:
            shutil.move(src, dest)
            result["moved"].append(dest)
        except OSError as ex:
            try:
                shutil.copy2(src, dest)
                archived_names.add(name)
                result["copied"].append(dest)
                result["failed"].append(f"{name}: copied to archive but source is still locked ({ex})")
            except OSError as copy_ex:
                result["failed"].append(f"{name}: {copy_ex}")
    if archived_names:
            write_archived_source_names(archived_names)
    return result

def load_existing_index():
    path = os.path.join(DATA_DIR, "codexes.json")
    if not os.path.exists(path):
        return []
    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []

def codex_summary_from_file(path, meta_keys):
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    if "entries" not in data:
        return None
    return {k: data.get(k) for k in meta_keys}, data

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--archive-sources",
        action="store_true",
        help="Move successfully converted source .docx files to 法典源_已转换/<timestamp>/ so future conversions skip them.",
    )
    args = parser.parse_args()

    archived_names = load_archived_source_names()
    docs = []
    for d in sorted(glob.glob(os.path.join(SRC_DIR, "*.docx"))):
        name = os.path.basename(d)
        if name in archived_names:
            print(f"[ARCHIVE-SKIP] {name}: already archived")
            continue
        docs.append(d)
    META_KEYS = ("id", "title", "version", "author", "entryCount", "imagedCount")
    existing_index = load_existing_index()
    produced_infos = []
    seen = {}
    produced = set()
    for d in docs:
        if os.path.basename(d).startswith("~$"):
            continue
        stem = os.path.splitext(os.path.basename(d))[0]
        cid = codex_id(stem)
        if cid in seen:                       # 防撞：同 id 自动加后缀，杜绝互相覆盖
            seen[cid] += 1
            cid = f"{cid}_{seen[cid]}"
        else:
            seen[cid] = 1
        info = convert(d, cid)
        produced.add(info["id"])
        produced_infos.append({k: info[k] for k in META_KEYS})
        print(f"[OK] {info['id']}: {info['entryCount']} entries, "
              f"{info['imagedCount']} imaged, {info['reviewCount']} to review")

    # 冻结保留：有数据文件、但本次没有对应 docx 的法典，原样保留在索引里。
    # 索引顺序优先沿用既有 codexes.json，避免清空 法典源/ 后重排顶部切换。
    produced_by_id = {item["id"]: item for item in produced_infos}
    kept_by_id = {}
    for jf in sorted(glob.glob(os.path.join(DATA_DIR, "*.json"))):
        cid = os.path.splitext(os.path.basename(jf))[0]
        if cid == "codexes" or cid in produced:
            continue
        try:
            summary, kept = codex_summary_from_file(jf, META_KEYS)
            if summary is None:         # 不是法典数据文件，跳过
                continue
            kept_by_id[cid] = (summary, kept)
        except Exception as ex:
            print(f"[SKIP] {os.path.basename(jf)}: {ex}")

    index = []
    emitted = set()
    for old in existing_index:
        cid = old.get("id")
        if cid in produced_by_id:
            index.append(produced_by_id[cid])
            emitted.add(cid)
        elif cid in kept_by_id:
            summary, kept = kept_by_id[cid]
            index.append(summary)
            emitted.add(cid)
            print(f"[KEEP] {cid}: kept from existing data (no docx) - {kept.get('entryCount')} entries")

    for item in produced_infos:
        cid = item["id"]
        if cid not in emitted:
            index.append(item)
            emitted.add(cid)

    for cid in sorted(kept_by_id):
        if cid in emitted:
            continue
        summary, kept = kept_by_id[cid]
        index.append(summary)
        emitted.add(cid)
        print(f"[KEEP] {cid}: kept from existing data (no docx) - {kept.get('entryCount')} entries")

    with io.open(os.path.join(DATA_DIR, "codexes.json"), "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False, indent=2)
    print(f"[DONE] {len(index)} codex(es) -> site/data/")

    if args.archive_sources:
        archived = archive_sources(docs)
        count = len(archived["moved"]) + len(archived["copied"])
        if count:
            first = (archived["moved"] or archived["copied"])[0]
            print(f"[ARCHIVE] archived {count} source docx file(s) -> {os.path.relpath(os.path.dirname(first), ROOT)}")
            for issue in archived["failed"]:
                print(f"[ARCHIVE-WARN] {issue}")
        else:
            print("[ARCHIVE] no source docx files to move")

if __name__ == "__main__":
    main()
